from flask import Flask, render_template, request, send_file
from docx import Document

app = Flask(__name__)

@app.route('/')
@app.route('/report')
def report():
    return render_template('report.html')

@app.route('/generate-report', methods=['POST'])

def generate_report():
    date = request.form['date']
    academicyear = request.form['academicyear']
    semester = request.form['semester']
    nameofevent = request.form['nameofevent']
    dateandtime = request.form['dateandtime']
    eventvenue = request.form['eventvenue']
    organizedby = request.form['organizedby'] 
    targetaudience = request.form['targetaudience']
    resourceperson = request.form['resourceperson']
    eventcontents = request.form['eventcontents'] 
    detailsofevent = request.form['detailsofevent']

    # Load template
    doc = Document('reportgentemplate.docx')

    # Replace fields in the template
    for p in doc.paragraphs:
        if 't1' in p.text:
            p.text = p.text.replace('t1', date)
        if 't2' in p.text:
            p.text = p.text.replace('t2', academicyear)
            p.runs[0].bold = True
        if 't3' in p.text:
            p.text = p.text.replace('t3', semester)
            p.runs[0].bold = True
        if 't4' in p.text:
            p.text = p.text.replace('t4', nameofevent)
        if 'tf1' in p.text:
            p.text = p.text.replace('tf1', eventcontents)
        if 'r1' in p.text:
            p.text = p.text.replace('r1', detailsofevent)
        
        for table in doc.tables:
         for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                if 't5' in cell_text:
                    cell.text = cell_text.replace('t5', dateandtime)
                if 't6' in cell_text:
                    cell.text = cell_text.replace('t6', eventvenue) 
                if 't7' in cell_text:
                    cell.text = cell_text.replace('t7', organizedby) 
                if 't8' in cell_text:
                    cell.text = cell_text.replace('t8', targetaudience)
                if 't9' in cell_text:
                    cell.text = cell_text.replace('t9', resourceperson)
                    
    # Save the Word document
    doc.save('report.docx')
    # Send the file as a response
    return send_file('report.docx', as_attachment=True)
if __name__ == '__main__':
    app.run(debug = True)